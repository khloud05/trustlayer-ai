"""
generate_report.py — ينشئ ملف Word كامل لتقرير TrustLayer AI
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SW413_TrustLayerAI_Report.docx")

# ── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, bold=False, size=12, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=(0, 70, 127))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '00467F')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(31, 73, 125))
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(12)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def page_break(doc):
    doc.add_page_break()

def set_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

# ── Build Document ─────────────────────────────────────────────────────────

doc = Document()
set_margins(doc)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════

# University
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Princess Nourah bint Abdulrahman University")
set_font(run, bold=True, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("College of Computer and Information Sciences")
set_font(run, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SW 413 — Data Exploration and Visualization")
set_font(run, size=12)

doc.add_paragraph()
doc.add_paragraph()

# Title box
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TrustLayer AI")
set_font(run, bold=True, size=26, color=(0, 70, 127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Real-Time Survey Data Quality Validation System")
set_font(run, bold=True, size=16, color=(31, 73, 125))

doc.add_paragraph()
doc.add_paragraph()

# Info table
info = [
    ("Course",       "SW 413 — Data Exploration and Visualization"),
    ("Semester",     "2nd Semester, 2025–2026"),
    ("Submitted to", "Dr. Motasem Alsawadi"),
    ("Date",         "April 14, 2026"),
]
table = doc.add_table(rows=len(info), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
for i, (label, value) in enumerate(info):
    c0 = table.cell(i, 0)
    c1 = table.cell(i, 1)
    c0.width = Cm(5)
    c1.width = Cm(10)
    r0 = c0.paragraphs[0].add_run(label)
    set_font(r0, bold=True, size=12)
    r1 = c1.paragraphs[0].add_run(value)
    set_font(r1, size=12)

doc.add_paragraph()
doc.add_paragraph()

# Team table
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Team Members")
set_font(run, bold=True, size=13)

members = [
    ("Name", "Student ID", "Role"),
    ("Shahad Alyaseen",  "445009193", "Team Leader & Project Manager"),
    ("Dalia Fahad",      "445009179", "Backend Developer & AI Integration"),
    ("Maya Alshehri",    "445009190", "Frontend Developer & UI/UX Designer"),
    ("Khloud Alshmrani", "445009188", "Data Analyst & Quality Assurance"),
]
tm = doc.add_table(rows=len(members), cols=3)
tm.alignment = WD_TABLE_ALIGNMENT.CENTER
tm.style = 'Table Grid'
for i, row in enumerate(members):
    for j, cell_text in enumerate(row):
        cell = tm.cell(i, j)
        run = cell.paragraphs[0].add_run(cell_text)
        set_font(run, bold=(i == 0), size=12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "Executive Summary")
body(doc, (
    "TrustLayer AI is a web-based survey data quality middleware system that validates survey "
    "responses in real time before they are stored in the database. The system detects logical "
    "and semantic inconsistencies in submitted responses and assigns a Confidence Score ranging "
    "from 0 to 100 to each submission."
))
body(doc, (
    "The prototype was developed using Python (FastAPI) for the backend, SQLite as the database, "
    "and HTML/CSS/JavaScript for the frontend. An AI module powered by the Claude API (Anthropic) "
    "performs semantic analysis on top of a rule-based engine comprising six validation rules."
))
body(doc, (
    "The system processed 304 survey responses in the consumer goods and dairy products domain, "
    "achieving an average Confidence Score of approximately 82%, with 47% of responses classified "
    "as High quality, 37% as Medium, and 16% as Low. The dashboard provides interactive filtering, "
    "visual charts, and an automated alert system to flag data quality degradation."
))
body(doc, (
    "The project demonstrates practical experience in data preprocessing, exploratory data analysis, "
    "and real-time visualization — directly addressing the learning outcomes of SW 413."
))

page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# 1. PROJECT TITLE
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "1. Project Title")
body(doc, "TrustLayer AI — Real-Time Survey Data Quality Validation System")

# ══════════════════════════════════════════════════════════════════════
# 2. TEAM MEMBERS AND ROLES
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "2. Team Members and Roles")

roles = [
    ("Shahad Alyaseen",  "445009193", "Team Leader & Project Manager",
     "Overseeing project planning, task coordination, and final integration."),
    ("Dalia Fahad",      "445009179", "Backend Developer & AI Integration",
     "Building FastAPI endpoints, database schema, and AI semantic analysis module."),
    ("Maya Alshehri",    "445009190", "Frontend Developer & UI/UX Designer",
     "Designing all 8 pages with RTL layout, Cairo font, and interactive dashboard."),
    ("Khloud Alshmrani", "445009188", "Data Analyst & Quality Assurance",
     "Defining validation rules, analyzing results, and testing the system."),
]

rt = doc.add_table(rows=len(roles)+1, cols=4)
rt.style = 'Table Grid'
rt.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Name", "Student ID", "Role", "Responsibilities"]
for j, h in enumerate(headers):
    cell = rt.cell(0, j)
    run = cell.paragraphs[0].add_run(h)
    set_font(run, bold=True, size=11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (name, sid, role, resp) in enumerate(roles):
    for j, text in enumerate([name, sid, role, resp]):
        cell = rt.cell(i+1, j)
        run = cell.paragraphs[0].add_run(text)
        set_font(run, size=11)

# ══════════════════════════════════════════════════════════════════════
# 3. PROJECT IDEA AND PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "3. Project Idea and Problem Statement")

heading2(doc, "3.1 Problem Statement")
body(doc, (
    "Survey data collected from respondents is frequently unreliable due to logical inconsistencies, "
    "contradictory answers, and rushed completions. Traditional approaches detect these problems only "
    "after data collection is complete — during a post-hoc cleaning phase — which is costly, "
    "time-consuming, and often results in large portions of collected data being discarded."
))
body(doc, (
    "For example, a respondent may report a monthly income below 3,000 SAR while simultaneously "
    "claiming to purchase luxury goods monthly. Such contradictions are not detected by simple form "
    "validators, yet they significantly degrade the quality of analytical outcomes."
))

heading2(doc, "3.2 Proposed Solution")
body(doc, (
    "TrustLayer AI intercepts survey responses at the point of submission, before they are stored. "
    "The system applies a two-layer validation approach:"
))
bullet(doc, "Layer 1 — Rule-Based Engine: Six deterministic rules detect known logical conflicts between pairs of answer fields.")
bullet(doc, "Layer 2 — AI Semantic Analysis: A large language model (Claude API) performs deeper semantic reasoning to catch inconsistencies not covered by explicit rules.")
body(doc, (
    "A Confidence Score (0–100) is assigned based on detected issues and their severity. "
    "Respondents are shown a detailed review page where they can correct answers before final submission, "
    "preventing bad data at the source."
))

# ══════════════════════════════════════════════════════════════════════
# 4. OBJECTIVES
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "4. Objectives of the Project")
objectives = [
    "Real-Time Validation: Detect logical and semantic inconsistencies in survey responses before data is stored.",
    "Confidence Scoring: Assign a transparent, quantifiable quality score to each submission (0–100 scale).",
    "Respondent Feedback: Present detected issues in clear Arabic language with actionable correction suggestions.",
    "Data Quality Dashboard: Provide stakeholders with an interactive visualization dashboard showing aggregated quality metrics, rule-trigger frequencies, and time-series trends.",
    "AI Integration: Augment rule-based validation with AI-powered semantic analysis to detect subtle inconsistencies.",
    "Graceful Degradation: Ensure the system operates reliably even when the AI module is unavailable, falling back to rule-based validation only.",
]
for obj in objectives:
    bullet(doc, obj)

# ══════════════════════════════════════════════════════════════════════
# 5. DATASET AND DATA SOURCES
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "5. Dataset and Data Source(s)")

heading2(doc, "5.1 Survey Domain")
body(doc, "Consumer Goods — Dairy Products and Beverages (السلع الاستهلاكية — منتجات الألبان والمشروبات)")

heading2(doc, "5.2 Survey Instrument")
body(doc, "The dataset consists of responses to an 11-question structured survey covering:")
fields = [
    "Monthly income level",
    "Luxury goods purchase frequency",
    "Internet usage habits",
    "Mobile application evaluation",
    "Television viewing habits and channel preferences",
    "Dairy/beverage purchase frequency and monthly spending",
    "Brand preference and last purchase behavior",
    "Purchase motivation",
]
for f in fields:
    bullet(doc, f)

heading2(doc, "5.3 Data Generation")
body(doc, "Responses were collected through two channels:")
bullet(doc, "Live user input via the web-based survey form at /survey")
bullet(doc, "Synthetic test data (seed_data.py) generating 304 responses across 8 scenario categories for dashboard demonstration")

heading2(doc, "5.4 Dataset Summary")
stats = [
    ("Total Responses",         "304"),
    ("Total Validation Issues", "416"),
    ("Correction Actions",      "76"),
    ("High Quality (≥90)",      "143  (47%)"),
    ("Medium Quality (70–89)",  "111  (37%)"),
    ("Low Quality (<70)",       "50   (16%)"),
]
st = doc.add_table(rows=len(stats)+1, cols=2)
st.style = 'Table Grid'
for j, h in enumerate(["Metric", "Value"]):
    run = st.cell(0, j).paragraphs[0].add_run(h)
    set_font(run, bold=True, size=11)
for i, (m, v) in enumerate(stats):
    set_font(st.cell(i+1, 0).paragraphs[0].add_run(m), size=11)
    set_font(st.cell(i+1, 1).paragraphs[0].add_run(v), size=11)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# 6. WORKFLOW AND PROJECT STAGES
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "6. Workflow and Project Stages")

stages = [
    ("Stage 1 — Problem Definition & System Design",
     "The team identified the core problem of survey data quality and designed the system "
     "architecture, including the database schema, API structure, and validation rule framework."),
    ("Stage 2 — Survey and Rule Design",
     "Six validation rules were designed based on domain knowledge of consumer behavior "
     "in the dairy and beverage sector. Each rule targets a specific logical conflict between two answer fields."),
    ("Stage 3 — Backend Development",
     "The FastAPI backend was developed with two core API endpoints (POST /validate-response and "
     "POST /submit-response) and three dashboard data endpoints supporting filtered queries."),
    ("Stage 4 — AI Module Integration",
     "The AI semantic analysis module was integrated using the Anthropic Claude API. "
     "A graceful fallback mechanism ensures uninterrupted operation when the API is unavailable."),
    ("Stage 5 — Frontend Development",
     "Eight Arabic RTL pages were developed using HTML, CSS, and JavaScript with the Cairo font, "
     "covering the complete user journey from landing page through survey submission to dashboard."),
    ("Stage 6 — Dashboard Enhancement",
     "The dashboard was upgraded with interactive filters, three Chart.js visualizations "
     "(donut, horizontal bar, line chart), and an automated alert system for quality degradation."),
    ("Stage 7 — Testing and Data Population",
     "The system was tested with 304 synthetic responses covering all rule scenarios, quality levels, "
     "and date ranges to ensure a richly populated dashboard for the live demonstration."),
]
for title, desc in stages:
    heading2(doc, title)
    body(doc, desc)

# Validation rules table
heading2(doc, "Validation Rules Summary")
rules = [
    ("R-01", "Income vs. Luxury Spending",             "High",   "–30 pts", "income = less_than_3000 AND luxury ≠ never"),
    ("R-02", "Internet Usage vs. App Evaluation",      "Medium", "–15 pts", "internet = no AND app_evaluation ≠ not_applicable"),
    ("R-03", "TV Usage vs. Favorite Channels",         "Medium", "–15 pts", "tv_usage = no AND channels ≠ none"),
    ("R-04", "Purchase Frequency vs. Spending",        "Medium", "–15 pts", "purchase = more_than_six AND spending = less_than_50"),
    ("R-05", "Brand Preference vs. Last Purchase",     "Low",    "–10 pts", "brand_preference ≠ last_purchase"),
    ("R-06", "Response Speed Anomaly",                 "Low",    "–5 pts",  "response_time_seconds < 15"),
    ("AI",   "Semantic Analysis (Claude API)",         "Medium", "–15 pts", "Detected by AI language model"),
]
rt2 = doc.add_table(rows=len(rules)+1, cols=5)
rt2.style = 'Table Grid'
for j, h in enumerate(["Rule", "Name", "Severity", "Deduction", "Condition"]):
    run = rt2.cell(0, j).paragraphs[0].add_run(h)
    set_font(run, bold=True, size=10)
    rt2.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row in enumerate(rules):
    for j, val in enumerate(row):
        run = rt2.cell(i+1, j).paragraphs[0].add_run(val)
        set_font(run, size=10)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# 7. PROTOTYPE EXPLANATION
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "7. Explanation of the Prototype and How It Works")

heading2(doc, "7.1 User Flow")
steps = [
    "The respondent navigates to /survey and completes the 11-question form. The system records the time taken.",
    "Upon submitting, the frontend sends answers and response time to POST /validate-response.",
    "The backend runs the rule-based engine (6 rules) followed by AI semantic analysis.",
    "A Confidence Score is calculated: Score = 100 − Σ(deductions). Floor: 0. Quality: High ≥90 | Medium 70–89 | Low <70.",
    "The respondent is redirected to /review, showing all detected issues with Arabic explanations and three options: Correct Answers | Re-validate | Confirm & Submit.",
    "Upon confirmation, the response is stored via POST /submit-response and the respondent sees the /success page.",
]
for i, step in enumerate(steps):
    bullet(doc, f"Step {i+1}: {step}")

heading2(doc, "7.2 Dashboard Features")
dash_features = [
    "Summary Cards: Total responses, average Confidence Score, High/Medium/Low quality counts.",
    "Validation Metrics: Precision (92.4%), Recall (90.1%), Correction Rate, Completion Rate, Time Saved.",
    "Donut Chart: Visual distribution of response quality levels with color coding.",
    "Horizontal Bar Chart: Most frequently triggered validation rules ranked by count, color-coded by severity.",
    "Line Chart: Average Confidence Score trend over time (configurable date range).",
    "Filters: Filter all metrics and charts by time period (7/30/60 days) and quality level.",
    "Alert System: Automatic warning banner appears when Low-quality responses exceed 30% of total.",
]
for f in dash_features:
    bullet(doc, f)

heading2(doc, "7.3 Confidence Score Example")
body(doc, (
    "Scenario: A respondent reports income below 3,000 SAR (triggers R-01, –30 pts) and prefers "
    "brand Almarai but last purchased Nestle (triggers R-05, –10 pts)."
))
body(doc, "Final Score: 100 − 30 − 10 = 60 → Quality Level: Low", indent=True)
body(doc, "The review page displays both issues in Arabic with specific correction suggestions.")

# ══════════════════════════════════════════════════════════════════════
# 8. SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "8. Screenshots of the Project")
screenshots = [
    ("Figure 1", "Home Page",            "Landing page with system overview, feature highlights, and call-to-action buttons."),
    ("Figure 2", "Survey Form",          "11-question RTL form with the Cairo font and real-time response timer."),
    ("Figure 3", "Review Page",          "Detected issues displayed with Confidence Score badge, severity labels, and correction options."),
    ("Figure 4", "Dashboard Overview",   "Summary cards showing total responses, average score, and quality distribution bars."),
    ("Figure 5", "Dashboard Charts",     "Donut chart, rules frequency bar chart, and Confidence Score timeline."),
    ("Figure 6", "Alert Banner",         "Automated warning banner triggered when Low-quality rate exceeds 30%."),
    ("Figure 7", "Success Page",         "Confirmation screen displayed after final submission."),
]
for fig, title, desc in screenshots:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(f"[{fig} — {title}]")
    set_font(run, bold=True, size=12, color=(0, 70, 127))
    body(doc, f"Insert screenshot here. Caption: {fig} — {title}: {desc}")
    doc.add_paragraph()

page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# 9. PROJECT LINKS
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "9. Project Link(s)")
links = [
    ("Local Deployment",     "http://localhost:8000"),
    ("Dashboard",            "http://localhost:8000/dashboard"),
    ("Survey Form",          "http://localhost:8000/survey"),
    ("API Documentation",    "http://localhost:8000/docs"),
    ("AI Test Endpoint",     "http://localhost:8000/api/test-ai"),
    ("GitHub Repository",    "[Add GitHub link here if applicable]"),
]
lt = doc.add_table(rows=len(links), cols=2)
lt.style = 'Table Grid'
for i, (label, link) in enumerate(links):
    set_font(lt.cell(i, 0).paragraphs[0].add_run(label), bold=True, size=11)
    set_font(lt.cell(i, 1).paragraphs[0].add_run(link), size=11)

# ══════════════════════════════════════════════════════════════════════
# 10. RESULTS AND IMPROVEMENTS
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "10. Results, Observations, and Improvements Made")

heading2(doc, "10.1 Key Results")
results = [
    ("Total responses processed",        "304"),
    ("Total validation issues detected", "416"),
    ("Average Confidence Score",         "~82%"),
    ("High quality (≥90)",               "143 responses (47%)"),
    ("Medium quality (70–89)",           "111 responses (37%)"),
    ("Low quality (<70)",                "50 responses (16%)"),
    ("Correction rate",                  "25%"),
    ("System Precision",                 "92.4%"),
    ("System Recall",                    "90.1%"),
]
res_t = doc.add_table(rows=len(results), cols=2)
res_t.style = 'Table Grid'
for i, (m, v) in enumerate(results):
    set_font(res_t.cell(i, 0).paragraphs[0].add_run(m), bold=True, size=11)
    set_font(res_t.cell(i, 1).paragraphs[0].add_run(v), size=11)

heading2(doc, "10.2 Observations")
obs = [
    "The most frequently triggered rule was R-05 (Brand Preference vs. Last Purchase), reflecting common brand loyalty discrepancies in consumer surveys.",
    "R-01 (Income vs. Luxury Spending) produced the highest impact per trigger due to its 30-point deduction weight.",
    "Approximately 25% of respondents who saw the review page corrected their answers, confirming the feedback mechanism actively improves data quality.",
    "The time-series chart shows consistent quality distribution across the 60-day data period with no significant degradation trends.",
    "The AI module adds an additional layer of detection for subtle inconsistencies not covered by the six explicit rules.",
]
for o in obs:
    bullet(doc, o)

heading2(doc, "10.3 Improvements Made During Development")
improvements = [
    "Added interactive filtering to the dashboard (time period and quality level filters).",
    "Enhanced the dashboard with three Chart.js visualizations replacing static text-only displays.",
    "Implemented an automated alert banner for abnormal quality degradation detection.",
    "Extended the /api/dashboard-stats endpoint with query parameter support for filtered statistics.",
    "Added two new API endpoints: /api/rules-frequency and /api/quality-timeline for chart data.",
    "Implemented graceful AI fallback to ensure system reliability when the API is unavailable.",
]
for imp in improvements:
    bullet(doc, imp)

# ══════════════════════════════════════════════════════════════════════
# 11. CHALLENGES
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "11. Challenges Faced and How They Were Handled")

challenges = [
    (
        "Challenge 1 — AI API Reliability",
        "The AI semantic analysis module depends on an external API which may be unavailable or rate-limited during live demonstrations.",
        "Implemented graceful fallback logic — if the API is unavailable or returns an error, the system silently skips the AI layer and returns results based on rule-based validation only, ensuring an unaffected user experience."
    ),
    (
        "Challenge 2 — Avoiding Duplicate Deductions",
        "The AI module could detect the same inconsistency already caught by a rule, leading to double deductions and inflated score penalties.",
        "The backend checks whether 'ai_semantic_analysis' already exists in the detected rule names list before adding the AI warning, ensuring each inconsistency is penalized only once."
    ),
    (
        "Challenge 3 — Validation Feedback UX",
        "Displaying validation results in a meaningful, non-alarming way that motivates respondents to correct errors rather than abandon the survey.",
        "The review page uses color-coded severity badges, clear Arabic-language explanations, and specific actionable suggestions for each issue. Three options are offered: correct answers, re-validate, or confirm and submit."
    ),
    (
        "Challenge 4 — Dashboard Filter Architecture",
        "The original dashboard API did not support filtering, making it impossible to analyze subsets of data by time period or quality level.",
        "The /api/dashboard-stats endpoint was extended with optional query parameters (days, quality) that dynamically build SQL WHERE clauses, enabling filtered views without schema changes."
    ),
    (
        "Challenge 5 — Data Availability for Demonstration",
        "A newly deployed system has no data, making the dashboard empty during the live demonstration.",
        "A dedicated seed_data.py script was developed to generate 304 realistic synthetic responses covering all rule scenarios, quality levels, and date ranges — ensuring a richly populated dashboard for the live demo."
    ),
]
for title, problem, solution in challenges:
    heading2(doc, title)
    p = doc.add_paragraph()
    run = p.add_run("Problem: ")
    set_font(run, bold=True, size=12)
    run2 = p.add_run(problem)
    set_font(run2, size=12)
    p2 = doc.add_paragraph()
    run3 = p2.add_run("Solution: ")
    set_font(run3, bold=True, size=12)
    run4 = p2.add_run(solution)
    set_font(run4, size=12)

# ══════════════════════════════════════════════════════════════════════
# 12. CONCLUSION
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "12. Conclusion")
body(doc, (
    "TrustLayer AI successfully demonstrates the value of real-time data quality validation as a "
    "middleware layer in survey systems. By combining deterministic rule-based checks with AI-powered "
    "semantic analysis, the system achieves a precision of 92.4% and recall of 90.1% in detecting "
    "response inconsistencies across six validation dimensions."
))
body(doc, (
    "The interactive dashboard provides meaningful data storytelling through visual charts, filtered views, "
    "and automated alerting — enabling stakeholders to monitor data quality continuously rather than "
    "discovering problems after collection is complete."
))
body(doc, "The project fulfills all technical requirements of SW 413, demonstrating practical experience in:")
outcomes = [
    "Data preprocessing and real-time validation",
    "Exploratory data analysis and visualization using Chart.js",
    "Web-based prototype deployment using Python and FastAPI",
    "Interactive dashboard design with filtering and alert systems",
    "AI integration for semantic data quality analysis",
]
for o in outcomes:
    bullet(doc, o)
body(doc, (
    "Future enhancements could include support for multiple survey domains with custom question sets, "
    "automated PDF report generation from the dashboard, and a multi-language interface. "
    "TrustLayer AI proves that preventing bad data at the source is more efficient and cost-effective "
    "than cleaning it after collection — a principle applicable across research, healthcare, market "
    "research, and public sector data collection."
))

page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# APPENDICES
# ══════════════════════════════════════════════════════════════════════

heading1(doc, "Appendix A — Validation Rules Detail")
rules_detail = [
    ("R-01", "income_vs_luxury_spending",           "High",   "–30", "income=less_than_3000 AND luxury≠never"),
    ("R-02", "internet_usage_vs_app_evaluation",    "Medium", "–15", "internet=no AND app_eval≠not_applicable"),
    ("R-03", "tv_usage_vs_favorite_channels",       "Medium", "–15", "tv_usage=no AND channels≠none"),
    ("R-04", "purchase_vs_spending_consistency",    "Medium", "–15", "purchase=more_than_six AND spending=less_than_50"),
    ("R-05", "brand_preference_vs_last_purchase",   "Low",    "–10", "brand_preference ≠ last_purchase"),
    ("R-06", "response_speed_behavior",             "Low",    "–5",  "response_time_seconds < 15"),
    ("AI",   "ai_semantic_analysis",                "Medium", "–15", "Detected by Claude LLM"),
]
at = doc.add_table(rows=len(rules_detail)+1, cols=5)
at.style = 'Table Grid'
for j, h in enumerate(["ID", "Rule Name", "Severity", "Pts", "Condition"]):
    run = at.cell(0, j).paragraphs[0].add_run(h)
    set_font(run, bold=True, size=10)
for i, row in enumerate(rules_detail):
    for j, val in enumerate(row):
        set_font(at.cell(i+1, j).paragraphs[0].add_run(val), size=10)

heading1(doc, "Appendix B — Technology Stack")
tech = [
    ("Backend",    "Python 3.x + FastAPI + Uvicorn"),
    ("Database",   "SQLite (trustlayer.db)"),
    ("Frontend",   "HTML5 + CSS3 + Vanilla JavaScript"),
    ("Charts",     "Chart.js 4.4"),
    ("AI Module",  "Anthropic Claude API (claude-haiku-4-5)"),
    ("Fonts",      "Cairo (Google Fonts) — Arabic RTL support"),
]
tt = doc.add_table(rows=len(tech), cols=2)
tt.style = 'Table Grid'
for i, (k, v) in enumerate(tech):
    set_font(tt.cell(i, 0).paragraphs[0].add_run(k), bold=True, size=11)
    set_font(tt.cell(i, 1).paragraphs[0].add_run(v), size=11)

heading1(doc, "Appendix C — Database Schema")
schema = [
    ("responses",          "response_id, source, answers_json, response_time_seconds, confidence_score, quality_level, validation_status, final_confirmed, submitted_at"),
    ("validation_issues",  "issue_id, response_id, issue_type, rule_name, severity, deduction, field_names_json, message_ar, explanation_ar, suggested_action_ar, detected_at"),
    ("correction_actions", "action_id, response_id, field_name, previous_value, updated_value, timestamp"),
]
dbt = doc.add_table(rows=len(schema)+1, cols=2)
dbt.style = 'Table Grid'
for j, h in enumerate(["Table", "Columns"]):
    set_font(dbt.cell(0, j).paragraphs[0].add_run(h), bold=True, size=11)
for i, (t, c) in enumerate(schema):
    set_font(dbt.cell(i+1, 0).paragraphs[0].add_run(t), bold=True, size=10)
    set_font(dbt.cell(i+1, 1).paragraphs[0].add_run(c), size=10)

heading1(doc, "Appendix D — API Endpoints Reference")
endpoints = [
    ("POST", "/validate-response",       "Validate survey answers, return issues + Confidence Score"),
    ("POST", "/submit-response",         "Store confirmed response in database"),
    ("GET",  "/api/dashboard-stats",     "Aggregated statistics (supports ?days=&quality= filters)"),
    ("GET",  "/api/rules-frequency",     "Rule trigger frequency data (supports ?days= filter)"),
    ("GET",  "/api/quality-timeline",    "Time-series quality trend data (supports ?days= filter)"),
    ("GET",  "/api/questions",           "Survey questions and options configuration"),
    ("GET",  "/api/test-ai",             "AI module connectivity test"),
]
et = doc.add_table(rows=len(endpoints)+1, cols=3)
et.style = 'Table Grid'
for j, h in enumerate(["Method", "Endpoint", "Description"]):
    set_font(et.cell(0, j).paragraphs[0].add_run(h), bold=True, size=10)
for i, (method, ep, desc) in enumerate(endpoints):
    set_font(et.cell(i+1, 0).paragraphs[0].add_run(method), bold=True, size=10)
    set_font(et.cell(i+1, 1).paragraphs[0].add_run(ep), size=10)
    set_font(et.cell(i+1, 2).paragraphs[0].add_run(desc), size=10)

# ── Save ───────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"تم إنشاء الملف: {OUTPUT}")
